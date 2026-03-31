from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt


def add_code_paragraph(doc: Document, code_line: str) -> None:
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(code_line)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # Global default font for normal text
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    in_code_block = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                doc.add_paragraph()
            continue

        if in_code_block:
            add_code_paragraph(doc, line)
            continue

        if not stripped:
            doc.add_paragraph()
            continue

        if stripped == "---":
            doc.add_paragraph("-" * 80)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            continue

        if stripped[:2].isdigit() and stripped[1] == ".":
            p = doc.add_paragraph(stripped, style="List Number")
            p.paragraph_format.space_after = Pt(3)
            continue

        p = doc.add_paragraph(stripped)
        p.paragraph_format.space_after = Pt(6)

    doc.save(docx_path)


if __name__ == "__main__":
    project_root = Path(__file__).resolve().parent
    if len(sys.argv) >= 2:
        md_file = Path(sys.argv[1]).resolve()
    else:
        md_file = project_root / "HuongDanChiTiet_DuAn_QuanLyPhongHop_2026.md"

    if len(sys.argv) >= 3:
        docx_file = Path(sys.argv[2]).resolve()
    else:
        docx_file = md_file.with_suffix(".docx")

    markdown_to_docx(md_file, docx_file)
    print(f"Created: {docx_file}")
